import os
import logging
from pathlib import Path
from typing import List, Dict, Optional, Union
import PyPDF2
import docx
from bs4 import BeautifulSoup
import pandas as pd
from config import Config

class DocumentProcessor:
    """Process various document types and extract text content"""
    
    def __init__(self, config: Config = None):
        self.config = config or Config()
        self.logger = logging.getLogger(__name__)
        
    def process_document(self, file_path: Union[str, Path]) -> str:
        """
        Process a document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.MAX_FILE_SIZE_MB:
            raise ValueError(f"File too large: {file_size_mb:.2f}MB > {self.config.MAX_FILE_SIZE_MB}MB")
        
        # Process based on file extension
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._process_pdf(file_path)
        elif extension == '.docx':
            return self._process_docx(file_path)
        elif extension in ['.txt', '.md']:
            return self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return self._clean_text(text)
        
        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
        
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
    
    def _process_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return self._clean_text(text)
        
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"Could not decode text file: {file_path}")
        
        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters that might cause issues
        text = text.replace('\x00', '')  # Remove null bytes
        text = text.replace('\r', '\n')  # Normalize line endings
        
        # Remove excessive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        Split text into chunks for processing
        
        Args:
            text: Input text
            chunk_size: Size of each chunk (default from config)
            overlap: Overlap between chunks (default from config)
            
        Returns:
            List of text chunks
        """
        chunk_size = chunk_size or self.config.CHUNK_SIZE
        overlap = overlap or self.config.CHUNK_OVERLAP
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings in the last part of the chunk
                last_sentence_end = text.rfind('.', start, end)
                if last_sentence_end != -1 and last_sentence_end > start + chunk_size // 2:
                    end = last_sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    def get_document_metadata(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing document metadata
        """
        file_path = Path(file_path)
        
        metadata = {
            'filename': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_type': file_path.suffix.lower(),
            'created_time': file_path.stat().st_ctime,
            'modified_time': file_path.stat().st_mtime,
        }
        
        # Extract additional metadata based on file type
        if file_path.suffix.lower() == '.pdf':
            metadata.update(self._get_pdf_metadata(file_path))
        elif file_path.suffix.lower() == '.docx':
            metadata.update(self._get_docx_metadata(file_path))
        
        return metadata
    
    def _get_pdf_metadata(self, file_path: Path) -> Dict[str, any]:
        """Extract metadata from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted,
                }
                
                # Extract PDF metadata if available
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_meta.get('/Title', ''),
                        'author': pdf_meta.get('/Author', ''),
                        'subject': pdf_meta.get('/Subject', ''),
                        'creator': pdf_meta.get('/Creator', ''),
                        'producer': pdf_meta.get('/Producer', ''),
                        'creation_date': pdf_meta.get('/CreationDate', ''),
                        'modification_date': pdf_meta.get('/ModDate', ''),
                    })
                
                return metadata
        
        except Exception as e:
            self.logger.error(f"Error extracting PDF metadata: {str(e)}")
            return {}
    
    def _get_docx_metadata(self, file_path: Path) -> Dict[str, any]:
        """Extract metadata from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            }
            
            # Extract core properties
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'category': core_props.category or '',
                'comments': core_props.comments or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision,
            })
            
            return metadata
        
        except Exception as e:
            self.logger.error(f"Error extracting DOCX metadata: {str(e)}")
            return {}
    
    def batch_process_documents(self, file_paths: List[Union[str, Path]]) -> List[Dict[str, any]]:
        """
        Process multiple documents in batch
        
        Args:
            file_paths: List of document file paths
            
        Returns:
            List of processing results
        """
        results = []
        
        for file_path in file_paths:
            try:
                text = self.process_document(file_path)
                metadata = self.get_document_metadata(file_path)
                
                result = {
                    'file_path': str(file_path),
                    'text': text,
                    'metadata': metadata,
                    'text_length': len(text),
                    'word_count': len(text.split()),
                    'status': 'success'
                }
                
            except Exception as e:
                result = {
                    'file_path': str(file_path),
                    'text': '',
                    'metadata': {},
                    'text_length': 0,
                    'word_count': 0,
                    'status': 'error',
                    'error_message': str(e)
                }
                
                self.logger.error(f"Error processing {file_path}: {str(e)}")
            
            results.append(result)
        
        return results
    
    def validate_document(self, file_path: Union[str, Path]) -> Dict[str, bool]:
        """
        Validate document before processing
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing validation results
        """
        file_path = Path(file_path)
        
        validation = {
            'exists': file_path.exists(),
            'readable': False,
            'supported_type': False,
            'size_ok': False,
            'not_empty': False
        }
        
        if validation['exists']:
            try:
                validation['readable'] = os.access(file_path, os.R_OK)
                validation['supported_type'] = file_path.suffix.lower() in self.config.SUPPORTED_FILE_TYPES
                
                file_size_mb = file_path.stat().st_size / (1024 * 1024)
                validation['size_ok'] = file_size_mb <= self.config.MAX_FILE_SIZE_MB
                validation['not_empty'] = file_path.stat().st_size > 0
                
            except Exception as e:
                self.logger.error(f"Error validating document {file_path}: {str(e)}")
        
        return validation
    
    def get_text_statistics(self, text: str) -> Dict[str, any]:
        """
        Generate statistics about the extracted text
        
        Args:
            text: Input text
            
        Returns:
            Dictionary containing text statistics
        """
        words = text.split()
        sentences = text.split('.')
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'average_words_per_sentence': len(words) / max(len(sentences), 1),
            'average_sentences_per_paragraph': len(sentences) / max(len(paragraphs), 1),
            'unique_words': len(set(word.lower() for word in words if word.isalpha())),
            'lexical_diversity': len(set(word.lower() for word in words if word.isalpha())) / max(len(words), 1)
        }